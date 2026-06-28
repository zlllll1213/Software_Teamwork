import csv
import io
import json
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Protocol

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


TEXT_EXTENSIONS = {
    ".adoc",
    ".bash",
    ".c",
    ".cc",
    ".cfg",
    ".cmake",
    ".conf",
    ".cpp",
    ".csv",
    ".cxx",
    ".fish",
    ".go",
    ".h",
    ".hpp",
    ".htm",
    ".html",
    ".ini",
    ".java",
    ".js",
    ".json",
    ".jsonl",
    ".jsx",
    ".kt",
    ".lua",
    ".md",
    ".markdown",
    ".php",
    ".pl",
    ".py",
    ".rb",
    ".rs",
    ".rst",
    ".sh",
    ".sql",
    ".toml",
    ".ts",
    ".tsx",
    ".tsv",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
    ".zsh",
}

KNOWN_TEXT_FILENAMES = {
    "AUTHORS",
    "CHANGELOG",
    "CHANGES",
    "CONTRIBUTING",
    "COPYING",
    "Dockerfile",
    "INSTALL",
    "LICENSE",
    "MAINTAINERS",
    "Makefile",
    "NEWS",
    "README",
}


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    parser_backend: str
    metadata: dict[str, object]


class Parser(Protocol):
    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        ...


class TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        stripped = data.strip()
        if stripped:
            self.parts.append(stripped)

    def text(self) -> str:
        return "\n".join(self.parts)


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def parse_text(filename: str, content: bytes) -> ParsedDocument:
    text = decode_text(content)
    suffix = Path(filename).suffix.lower()
    if suffix in {".htm", ".html", ".xml"}:
        extractor = TextExtractor()
        extractor.feed(text)
        text = extractor.text()
    elif suffix == ".json":
        try:
            text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            pass
    return ParsedDocument(text=text, parser_backend="text", metadata={"suffix": suffix})


def parse_pdf(filename: str, content: bytes) -> ParsedDocument:
    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"Page {index}\n{page_text}")
    return ParsedDocument(
        text="\n\n".join(pages),
        parser_backend="pypdf",
        metadata={"suffix": ".pdf", "page_count": len(reader.pages)},
    )


def parse_docx(filename: str, content: bytes) -> ParsedDocument:
    document = DocxDocument(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return ParsedDocument(
        text="\n\n".join(parts),
        parser_backend="python-docx",
        metadata={"suffix": ".docx"},
    )


def parse_xlsx(filename: str, content: bytes) -> ParsedDocument:
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append("\t".join(values))
    return ParsedDocument(
        text="\n".join(parts),
        parser_backend="openpyxl",
        metadata={"suffix": ".xlsx", "sheet_count": len(workbook.worksheets)},
    )


def parse_csv_or_tsv(filename: str, content: bytes) -> ParsedDocument:
    text = decode_text(content)
    suffix = Path(filename).suffix.lower()
    dialect = "excel-tab" if suffix == ".tsv" else "excel"
    rows: list[str] = []
    try:
        reader = csv.reader(io.StringIO(text), dialect=dialect)
        for row in reader:
            rows.append("\t".join(row))
    except csv.Error:
        rows = text.splitlines()
    return ParsedDocument(text="\n".join(rows), parser_backend="csv", metadata={"suffix": suffix})


def is_supported_filename(filename: str) -> bool:
    path = Path(filename)
    suffix = path.suffix.lower()
    return (
        suffix in TEXT_EXTENSIONS
        or suffix in {".pdf", ".docx", ".xlsx"}
        or path.name in KNOWN_TEXT_FILENAMES
    )


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    path = Path(filename)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(filename, content)
    if suffix == ".docx":
        return parse_docx(filename, content)
    if suffix == ".xlsx":
        return parse_xlsx(filename, content)
    if suffix in {".csv", ".tsv"}:
        return parse_csv_or_tsv(filename, content)
    if suffix in TEXT_EXTENSIONS or path.name in KNOWN_TEXT_FILENAMES:
        return parse_text(filename, content)
    raise ValueError(f"Unsupported file type: {filename}")
